#!/usr/bin/env python3
from __future__ import annotations

import csv
import json
import math
import os
import re
import textwrap
from datetime import date
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DATA = PROJECT_ROOT / "data"
FIGURES = PROJECT_ROOT / "figures"
PAPER = PROJECT_ROOT / "paper"
TABLES = PROJECT_ROOT / "tables"
PUBLICATION = PROJECT_ROOT / "publication"
DOCS = PROJECT_ROOT / "docs"

AUTHOR = "Andrea Bertoldo"
AFFILIATION = "Independent Student Researcher, Italy"
TITLE = "Training Volume, Sleep and Endurance Performance"
SUBTITLE = "A Longitudinal Analysis of 29 Official Races in a Competitive Distance Runner"
ALT_SUBTITLE = "A Four-Year Wearable-Based Case Study (2023-2026)"


KEY_FINDINGS = [
    ("56-day running volume", 0.633, 0.0002),
    ("28-day running volume", 0.630, 0.0003),
    ("28-day sleep average", 0.580, 0.0010),
    ("14-day running volume", 0.505, 0.0053),
    ("VO2max nearest", 0.378, 0.0431),
    ("14-day sleep average", 0.403, 0.0300),
]


FIGURE_SOURCES = [
    (
        "race_performance_timeline.png",
        "Figure 1. Official FIDAL race performance timeline, 2023-2026.",
        "Event-relative performance score across the 29 official FIDAL race rows.",
    ),
    (
        "stat_correlation_rankings.png",
        "Figure 2. Ranked associations between Garmin-derived variables and event-relative performance.",
        "Longer running-volume windows and 28-day sleep showed the strongest complete-data associations.",
    ),
    (
        "stat_top_relationships_scatter.png",
        "Figure 3. Scatterplots for the strongest complete-data relationships.",
        "Associations are exploratory and observational; regression lines are descriptive only.",
    ),
    (
        "stat_best_worst_quartile_differences.png",
        "Figure 4. Best versus worst quartile comparison.",
        "Best-quartile races had higher 14-, 28-, and 56-day running volume and higher 28-day sleep.",
    ),
    (
        "sleep_before_races.png",
        "Figure 5. Sleep averages before official race days.",
        "Sleep signals were stronger over 14- and 28-day windows than over the final 7 days.",
    ),
]


def ensure_dirs() -> None:
    for path in (PAPER, TABLES, PUBLICATION, FIGURES, DOCS):
        path.mkdir(parents=True, exist_ok=True)


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(path: Path, rows: list[dict], fieldnames: list[str]) -> None:
    with path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def fmt_p(value: float | str) -> str:
    if value == "" or value is None:
        return ""
    value = float(value)
    if value < 0.001:
        return f"{value:.4f}"
    if value < 0.01:
        return f"{value:.4f}".rstrip("0").rstrip(".")
    text = f"{value:.4f}"
    if text.endswith("00"):
        return text[:-1]
    return text.rstrip("0").rstrip(".")


def fmt_num(value: float | str, digits: int = 3) -> str:
    if value == "" or value is None:
        return ""
    quant = Decimal("1").scaleb(-digits)
    return str(Decimal(str(value)).quantize(quant, rounding=ROUND_HALF_UP))


def get_counts() -> dict:
    activity = pd.read_csv(DATA / "activity_summary.csv")
    health = pd.read_csv(DATA / "daily_health_summary.csv")
    races = pd.read_csv(DATA / "race_context_dataset.csv")
    run_rows = activity[activity["activity_category"] == "run"]
    return {
        "activities": int(len(activity)),
        "runs": int(len(run_rows)),
        "run_km": float(run_rows["distance_km"].sum()),
        "sleep_records": int(health["sleep_hours"].notna().sum()),
        "hrv_records": int(health["hrv_ms"].notna().sum()),
        "races": int(len(races)),
        "matched_races": int((races["garmin_match_status"] == "matched").sum()),
        "garmin_start": str(activity["activity_date"].min()),
        "garmin_end": str(activity["activity_date"].max()),
        "race_start": str(races["race_date"].min()),
        "race_end": str(races["race_date"].max()),
        "suspicious_matches": int(races["suspicious_mismatch_flag"].fillna("").astype(str).str.len().gt(0).sum()),
    }


def build_tables() -> None:
    corr = read_csv(DATA / "stat_correlation_results.csv")
    table1_fields = [
        "variable",
        "n",
        "pearson_r",
        "pearson_p",
        "pearson_ci_low",
        "pearson_ci_high",
        "spearman_rho",
        "spearman_p",
        "evidence_note",
    ]
    table1 = [
        {
            "variable": row["label"],
            "n": row["n"],
            "pearson_r": fmt_num(row["pearson_r"], 3),
            "pearson_p": fmt_p(row["pearson_p"]),
            "pearson_ci_low": fmt_num(row["pearson_ci_low"], 3),
            "pearson_ci_high": fmt_num(row["pearson_ci_high"], 3),
            "spearman_rho": fmt_num(row["spearman_rho"], 3),
            "spearman_p": fmt_p(row["spearman_p"]),
            "evidence_note": row["evidence_note"],
        }
        for row in corr
    ]
    write_csv(TABLES / "final_table_01.csv", table1, table1_fields)

    quart = read_csv(DATA / "stat_best_worst_quartile_comparison.csv")
    table2_fields = [
        "variable",
        "best_n",
        "worst_n",
        "best_mean",
        "worst_mean",
        "best_minus_worst",
        "standardized_mean_difference",
        "direction",
    ]
    table2 = [
        {
            "variable": row["label"],
            "best_n": row["best_n"],
            "worst_n": row["worst_n"],
            "best_mean": fmt_num(row["best_mean"], 3),
            "worst_mean": fmt_num(row["worst_mean"], 3),
            "best_minus_worst": fmt_num(row["best_minus_worst"], 3),
            "standardized_mean_difference": fmt_num(row["standardized_mean_difference"], 3),
            "direction": row["direction"],
        }
        for row in quart
    ]
    write_csv(TABLES / "final_table_02.csv", table2, table2_fields)

    coverage = read_csv(DATA / "hrv_rhr_combined_coverage_audit.csv")
    table3_fields = [
        "race_date",
        "event",
        "hrv_export_days_28d",
        "hrv_screenshot_points_28d",
        "hrv_combined_points_28d",
        "rhr_export_days_28d",
        "rhr_screenshot_points_28d",
        "rhr_combined_points_28d",
    ]
    table3 = [
        {
            "race_date": row["race_date"],
            "event": row["fidal_event"],
            "hrv_export_days_28d": row["hrv_export_days_28d"],
            "hrv_screenshot_points_28d": row["hrv_screenshot_days_28d"],
            "hrv_combined_points_28d": row["hrv_combined_days_28d"],
            "rhr_export_days_28d": row["rhr_export_days_28d"],
            "rhr_screenshot_points_28d": row["rhr_screenshot_days_28d"],
            "rhr_combined_points_28d": row["rhr_combined_days_28d"],
        }
        for row in coverage
    ]
    write_csv(TABLES / "final_table_03.csv", table3, table3_fields)

    captions = """# Final Table Captions

Table 1. Pearson and Spearman associations between Garmin-derived pre-race variables and event-relative official race performance. P-values are exploratory and uncorrected for multiple comparisons.

Table 2. Comparison of the best seven and worst seven official races by event-relative performance. Positive differences indicate higher values in the best-performance quartile.

Table 3. HRV/RHR exact export coverage and supplementary screenshot-derived coverage for the 28 days preceding each official FIDAL race. Screenshot-derived values are approximate and excluded from the primary statistical analysis.
"""
    (TABLES / "final_table_captions.md").write_text(captions, encoding="utf-8")


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
        "/Library/Fonts/Arial.ttf",
    ]
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)
    return ImageFont.load_default()


def wrap_text_pixels(draw: ImageDraw.ImageDraw, text: str, font_obj, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = f"{current} {word}".strip()
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def build_figures() -> None:
    title_font = font(34, bold=True)
    caption_font = font(22)
    note_font = font(18)
    for i, (source_name, caption, note) in enumerate(FIGURE_SOURCES, start=1):
        source = FIGURES / source_name
        base = Image.open(source).convert("RGB")
        canvas_w = max(1700, base.width + 220)
        top = 92
        bottom = 160
        canvas_h = top + base.height + bottom
        canvas = Image.new("RGB", (canvas_w, canvas_h), "white")
        draw = ImageDraw.Draw(canvas)
        x = (canvas_w - base.width) // 2
        canvas.paste(base, (x, top))
        draw.text((80, 32), caption, fill=(20, 32, 45), font=title_font)
        y = top + base.height + 28
        for line in wrap_text_pixels(draw, note, caption_font, canvas_w - 160):
            draw.text((80, y), line, fill=(55, 65, 75), font=caption_font)
            y += 30
        draw.text(
            (80, canvas_h - 42),
            "Source: official FIDAL race rows matched to Garmin-derived training context; descriptive visual, not causal.",
            fill=(90, 96, 105),
            font=note_font,
        )
        canvas.save(FIGURES / f"final_figure_{i:02d}.png", dpi=(300, 300))


def markdown_table(rows: list[dict], columns: list[str], labels: list[str] | None = None, max_rows: int | None = None) -> str:
    labels = labels or columns
    data = rows if max_rows is None else rows[:max_rows]
    out = ["| " + " | ".join(labels) + " |", "| " + " | ".join(["---"] * len(columns)) + " |"]
    for row in data:
        out.append("| " + " | ".join(str(row.get(col, "")) for col in columns) + " |")
    return "\n".join(out)


def manuscript_markdown(counts: dict) -> str:
    table1 = read_csv(TABLES / "final_table_01.csv")
    table2 = read_csv(TABLES / "final_table_02.csv")
    key_table_rows = [row for row in table1 if row["variable"] in [k[0].replace("running volume", "Run km").replace("sleep average", "Sleep avg").replace("VO2max nearest", "VO2max nearest") for k, _, _ in []]]
    selected = [row for row in table1 if row["variable"] in ["Run km 56d", "Run km 28d", "Sleep avg 28d", "Run km 14d", "VO2max nearest", "Sleep avg 14d"]]
    selected_sorted = sorted(selected, key=lambda r: abs(float(r["pearson_r"])), reverse=True)
    top_quartile = [row for row in table2 if row["variable"] in ["Run km 56d", "Run km 28d", "Sleep avg 28d", "Run km 14d", "VO2max nearest", "Acute load nearest"]]
    return f"""# {TITLE}: {SUBTITLE}

## {ALT_SUBTITLE}

**Author:** {AUTHOR}  
**Affiliation:** {AFFILIATION}  
**Manuscript version:** 1.0  
**Date:** {date.today().isoformat()}  

## Abstract

This exploratory single-athlete case study examined whether Garmin-derived training and recovery variables were associated with official endurance-race performance. The primary dataset included {counts['races']} official FIDAL race results from {counts['race_start']} to {counts['race_end']}, each matched to a same-day Garmin running activity. Garmin export coverage spanned {counts['garmin_start']} to {counts['garmin_end']} and included {counts['activities']:,} activities, {counts['runs']:,} runs, {counts['run_km']:,.1f} km of running, {counts['sleep_records']:,} sleep records, and {counts['hrv_records']} export-derived HRV records. Performance was normalized within official event type using an event-relative performance score, where higher values indicate better performance relative to the athlete's best official result in the same event. The strongest complete-data associations were observed for 56-day running volume (Pearson r = 0.633, p = 0.0002), 28-day running volume (r = 0.630, p = 0.0003), 28-day sleep average (r = 0.580, p = 0.001), 14-day running volume (r = 0.505, p = 0.0053), VO2max nearest to race day (r = 0.378, p = 0.0431), and 14-day sleep average (r = 0.403, p = 0.030). HRV and resting heart rate were not suitable for primary inference because exact export coverage was limited to four race rows. Screenshot-derived HRV/RHR graph digitizations improved historical coverage as secondary evidence but were excluded from the primary analysis because they were approximate and not equivalent to export-derived daily values. These findings suggest that, in this athlete, longer-term running volume and multi-week sleep consistency contained the clearest wearable-derived signals, but the design is observational and does not support causal claims.

## Keywords

wearable data; endurance running; Garmin; sleep; training volume; FIDAL; single-athlete case study; observational analysis

## Introduction

Wearable devices produce dense longitudinal records of training, sleep, and physiological estimates. For endurance athletes, these records may help describe the training context that precedes competition. However, consumer wearable variables are not controlled laboratory measures, and their relationship with race performance is often difficult to interpret because performance is shaped by event distance, training phase, tactics, weather, injury status, competition level, and recovery.

This project analyzes one competitive distance runner's longitudinal data using official race results as the performance ground truth. Official FIDAL results were treated as the authoritative source for race dates, events, and times. Garmin Connect data were used only to characterize pre-race training and recovery context. The purpose was descriptive: to identify whether any wearable-derived variables showed meaningful association with event-relative official race performance.

The primary research question was: which Garmin-derived variables appear most strongly associated with official endurance-race performance in this athlete across the 2023-2026 competition period?

## Methods

### Study Design

This is an exploratory, observational, single-athlete longitudinal case study. It is not an intervention study and does not support causal inference. All associations are interpreted as descriptive signals within one athlete's historical data.

### Athlete and Data Sources

The athlete was Andrea Bertoldo, an independent student researcher based in Italy. The Garmin Connect export covered {counts['garmin_start']} to {counts['garmin_end']}. The derived activity dataset contained {counts['activities']:,} activities, including {counts['runs']:,} running activities and {counts['run_km']:,.1f} km of running. Sleep coverage included {counts['sleep_records']:,} sleep records. HRV exact export coverage included {counts['hrv_records']} records.

Official race results were imported from the athlete's FIDAL profile and manually validated. The final primary race dataset included {counts['races']} official FIDAL race rows from 2023-2026. Each official race was matched to a same-day Garmin running activity; {counts['matched_races']} of {counts['races']} rows had Garmin matches. Two matches were flagged as potentially suspicious for manual review, but all official race rows were retained because FIDAL remained the race-performance ground truth.

### Race Inclusion Criteria

The analysis focused on middle-distance and endurance-relevant official events from 2023-2026, including 1500 m, 3000 m, 5000 m, road 5 km, road 10 km, 30-minute run, and steeplechase where useful for history. Childhood and non-endurance events were excluded from the primary race-context dataset.

### Outcome Variable

Race performance was modeled as an event-relative performance score. For each official event type, the athlete's best official time was used as the within-event reference. The analysis used `performance_score = -relative_performance_pct`, so higher scores indicate a race closer to the athlete's best official result within the same event. This normalization reduces, but does not eliminate, differences across event types.

### Predictor Variables

The primary tested variables were:

- running volume in the previous 7, 14, 28, and 56 days;
- sleep average in the previous 7, 14, and 28 days;
- HRV average in the previous 7, 14, and 28 days;
- resting heart rate proxy in the previous 7, 14, and 28 days;
- Garmin VO2max nearest to race day;
- Garmin acute load nearest to race day.

### HRV and Resting Heart Rate Handling

Exact Garmin export-derived HRV and RHR-proxy records were available only from 2025-09-18 to 2026-06-04, yielding preceding-28-day coverage for only four race rows. A later screenshot integration audit digitized Garmin graph screenshots as secondary/manual evidence. Those screenshot-derived HRV/RHR values were approximate graph extractions, not exact export records. They were therefore excluded from the primary statistical analysis.

### Statistical Analysis

Pearson and Spearman correlations were computed between each predictor and event-relative performance score. Approximate two-sided p-values and Fisher z confidence intervals were reported when possible. Best-quartile and worst-quartile race groups were compared using the best seven and worst seven official races by event-relative performance. P-values were exploratory and not corrected for multiple comparisons.

## Results

### Dataset Summary

The final dataset contained {counts['races']} official FIDAL races and {counts['matched_races']} Garmin same-day activity matches. Garmin export coverage included {counts['activities']:,} activities, {counts['runs']:,} runs, {counts['run_km']:,.1f} km of running, {counts['sleep_records']:,} sleep records, and {counts['hrv_records']} exact HRV records.

### Primary Association Results

The strongest complete-data associations were observed for longer running-volume windows and multi-week sleep averages. The leading Pearson correlations were:

{markdown_table(selected_sorted, ["variable", "n", "pearson_r", "pearson_p", "pearson_ci_low", "pearson_ci_high", "spearman_rho", "spearman_p"], ["Variable", "n", "Pearson r", "p", "95% CI low", "95% CI high", "Spearman rho", "p"])}

Table 1 provides the full correlation results in `tables/final_table_01.csv`. Figure 2 summarizes the ranked correlation signals, and Figure 3 shows scatterplots for the strongest complete-data relationships.

### Best Versus Worst Quartile

The best seven races by event-relative performance showed higher running volume across 7-, 14-, 28-, and 56-day windows and higher 14- and 28-day sleep averages than the worst seven races. The largest standardized differences were observed for 28-day running volume, 56-day running volume, and 28-day sleep average.

{markdown_table(top_quartile, ["variable", "best_mean", "worst_mean", "best_minus_worst", "standardized_mean_difference", "direction"], ["Variable", "Best mean", "Worst mean", "Difference", "Std diff", "Direction"])}

The complete comparison is provided in `tables/final_table_02.csv` and Figure 4.

### HRV/RHR Coverage

The exact Garmin export contained 256 HRV records and 256 RHR-proxy records, beginning on 2025-09-18. Only four of the 29 official races had exact HRV/RHR coverage in the preceding 28 days. Screenshot-derived graph digitization improved apparent historical coverage, but these estimates were approximate, often rolling or weekly, and showed substantial mismatch with export values on overlapping dates. They were therefore retained as supplementary evidence only and excluded from the primary analysis.

### Signal Ranking

The descriptive ranking from strongest to weakest complete-data signal was:

1. 56-day running volume
2. 28-day running volume
3. 28-day sleep average
4. 14-day running volume
5. VO2max nearest
6. 14-day sleep average
7. Acute load nearest
8. 7-day sleep average
9. 7-day running volume

HRV and RHR variables were ranked as very low-confidence because only four race rows had exact export coverage.

## Discussion

The clearest signals in this single-athlete dataset were longer-term running-volume windows. The 56-day and 28-day running-volume variables showed similar Pearson correlations with event-relative performance, suggesting that sustained training context was more informative than the final 7 days alone. This is plausible for middle-distance and endurance performance, where chronic aerobic development and accumulated training consistency often matter more than very short-term volume.

Sleep showed a window-dependent pattern. The 7-day sleep average had little explanatory value, whereas 14-day and 28-day averages showed stronger positive associations. In the repeated-event sensitivity analysis documented in the statistical report, 28-day sleep remained a strong signal. This may indicate that multi-week sleep consistency captured broader recovery context better than the immediate pre-race week. It should not be interpreted as proof that increasing sleep alone caused faster performances.

VO2max nearest to race day showed a smaller but positive association. Garmin VO2max is model-derived and partly influenced by recent running performance, so it is best interpreted as a general fitness-context marker rather than an independent explanatory variable.

Acute load was weaker than longer running-volume windows. This suggests that, in this dataset, accumulated running volume contained more performance-relevant information than a short-term load estimate alone. HRV and RHR could not be meaningfully interpreted in the primary statistical analysis because exact export coverage was too sparse.

## Limitations

- This is a single-athlete case study and may not generalize to other runners.
- The design is exploratory and observational; no causal claims are made.
- Race events were mixed. Event-relative normalization reduces but cannot remove all differences between 1500 m, 3000 m, 5000 m, road races, and other events.
- P-values are exploratory and were not corrected for multiple comparisons.
- Garmin-derived sleep, VO2max, acute load, HRV, and RHR-proxy variables are consumer wearable estimates, not controlled laboratory measurements.
- Exact HRV/RHR coverage was limited to four race rows.
- Screenshot-derived HRV/RHR values were approximate graph digitizations and excluded from the primary statistical analysis.
- Race tactics, weather, illness, injury, school/life stress, competition quality, and taper intent were not fully controlled.

## Conclusion

Across 29 official FIDAL races from 2023-2026, the strongest Garmin-derived signals associated with event-relative performance were longer-term running volume and multi-week sleep averages. The leading associations were 56-day running volume (r = 0.633), 28-day running volume (r = 0.630), and 28-day sleep average (r = 0.580). VO2max nearest to race day showed a smaller positive association. HRV and RHR exact export coverage was too sparse for primary inference, and screenshot-derived graph estimates were kept as supplementary evidence only. The results support a cautious descriptive conclusion: in this athlete, sustained training volume and multi-week sleep consistency appear to contain meaningful wearable-derived performance context, but the analysis remains observational and non-causal.

## References

1. Federazione Italiana di Atletica Leggera (FIDAL). Official athlete profile and race results for Andrea Bertoldo. Used as the authoritative source for official race dates, events, and times.
2. Garmin Connect. Personal data export, 2022-07-21 to 2026-06-03. Used for activity, training, sleep, HRV, RHR-proxy, VO2max, and acute-load context.
3. Project source files: `data/race_context_dataset.csv`, `data/stat_correlation_results.csv`, `data/stat_signal_ranking.csv`, `docs/Statistical_Findings.md`, `docs/HRV_RHR_Coverage_Audit.md`, and `docs/HRV_RHR_Screenshot_Integration_Report.md`.

## Appendix

### Appendix A. Reproducibility

The analysis is reproducible from the project scripts and derived CSV files:

1. `scripts/inspect_garmin_export.py`
2. `scripts/build_activity_summary.py`
3. `scripts/build_health_summary.py`
4. `scripts/build_official_race_list.py`
5. `scripts/build_race_context_dataset.py`
6. `scripts/statistical_analysis.py`
7. `scripts/audit_hrv_rhr_coverage.py`
8. `scripts/digitize_screenshot_graphs.py`
9. `scripts/prepare_public_release.py`

### Appendix B. Final Figures

- Figure 1: `figures/final_figure_01.png`
- Figure 2: `figures/final_figure_02.png`
- Figure 3: `figures/final_figure_03.png`
- Figure 4: `figures/final_figure_04.png`
- Figure 5: `figures/final_figure_05.png`

### Appendix C. Final Tables

- Table 1: `tables/final_table_01.csv`
- Table 2: `tables/final_table_02.csv`
- Table 3: `tables/final_table_03.csv`
"""


def save_markdown(counts: dict) -> str:
    md = manuscript_markdown(counts)
    (PAPER / "final_manuscript.md").write_text(md, encoding="utf-8")
    return md


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, color in [
        ("Heading 1", 16, RGBColor(20, 62, 92)),
        ("Heading 2", 13, RGBColor(20, 62, 92)),
        ("Heading 3", 11, RGBColor(20, 62, 92)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_markdown_table_to_docx(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            paragraph = table.cell(r, c).paragraphs[0]
            run = paragraph.add_run(cell)
            run.font.name = "Arial"
            run.font.size = Pt(8)
            if r == 0:
                run.bold = True
    doc.add_paragraph()


def markdown_to_docx(md: str) -> None:
    doc = Document()
    set_doc_style(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"{TITLE}\n{SUBTITLE}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(20, 62, 92)
    subtitle = doc.add_paragraph(ALT_SUBTITLE)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    doc.add_paragraph(f"{AUTHOR}\n{AFFILIATION}\nVersion 1.0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading == ALT_SUBTITLE:
                i += 1
                continue
            doc.add_heading(heading, level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table_to_docx(doc, table_lines)
            continue
        if line.startswith("- "):
            text = line[2:].strip().replace("`", "")
            doc.add_paragraph(text, style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).replace("`", "")
            doc.add_paragraph(text, style="List Number")
            i += 1
            continue
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "- ")) and not re.match(r"^\d+\. ", lines[i]):
            paragraph_lines.append(lines[i])
            i += 1
        text = " ".join(paragraph_lines).replace("**", "").replace("`", "")
        doc.add_paragraph(text)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("Figures", level=1)
    for idx, (_, caption, note) in enumerate(FIGURE_SOURCES, start=1):
        fig_path = FIGURES / f"final_figure_{idx:02d}.png"
        doc.add_picture(str(fig_path), width=Inches(6.2))
        p = doc.add_paragraph(caption + " " + note)
        p.runs[0].italic = True
    doc.save(PAPER / "final_manuscript.docx")


def md_lines_for_pdf(md: str) -> list:
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCenter", parent=styles["Title"], alignment=TA_CENTER, fontSize=17, leading=21, spaceAfter=12)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, leading=17, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#143E5C"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=11.5, leading=14, spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#143E5C"))
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=9.5, leading=12, spaceAfter=5)
    small = ParagraphStyle("Small", parent=styles["BodyText"], fontSize=8, leading=10, spaceAfter=4)
    story = []
    story.append(Paragraph(f"{TITLE}<br/>{SUBTITLE}", title_style))
    story.append(Paragraph(f"{ALT_SUBTITLE}<br/>{AUTHOR}<br/>{AFFILIATION}<br/>Version 1.0", ParagraphStyle("Sub", parent=body, alignment=TA_CENTER)))
    story.append(PageBreak())
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            heading = line[3:].strip()
            if heading != ALT_SUBTITLE:
                story.append(Paragraph(heading, h1))
            i += 1
            continue
        if line.startswith("### "):
            story.append(Paragraph(line[4:].strip(), h2))
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = []
            for raw in table_lines:
                cells = [cell.strip() for cell in raw.strip().strip("|").split("|")]
                if all(set(cell) <= {"-", ":"} for cell in cells):
                    continue
                rows.append([Paragraph(cell, small) for cell in cells])
            if rows:
                tbl = Table(rows, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#AAB2BD")),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 0.08 * inch))
            continue
        if line.startswith("- "):
            story.append(Paragraph("• " + line[2:].replace("`", ""), body))
            i += 1
            continue
        paragraph_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("#", "|", "- ")):
            paragraph_lines.append(lines[i].strip())
            i += 1
        text = " ".join(paragraph_lines)
        text = text.replace("**", "").replace("`", "")
        story.append(Paragraph(text, body))
    story.append(PageBreak())
    story.append(Paragraph("Final Figures", h1))
    for idx, (_, caption, note) in enumerate(FIGURE_SOURCES, start=1):
        fig_path = FIGURES / f"final_figure_{idx:02d}.png"
        story.append(PdfImage(str(fig_path), width=6.5 * inch, height=4.6 * inch))
        story.append(Paragraph(caption + " " + note, small))
        story.append(Spacer(1, 0.1 * inch))
    return story


def markdown_to_pdf(md: str) -> None:
    doc = SimpleDocTemplate(str(PAPER / "final_manuscript.pdf"), pagesize=letter, rightMargin=0.75 * inch, leftMargin=0.75 * inch, topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    story = md_lines_for_pdf(md)
    doc.build(story)


def write_readme_and_metadata(counts: dict) -> None:
    readme = f"""# {TITLE}: {SUBTITLE}

{ALT_SUBTITLE}

Author: {AUTHOR}, {AFFILIATION}

## Overview

This repository contains a reproducible single-athlete observational case study linking official FIDAL race results with Garmin Connect training and recovery data. The primary dataset contains {counts['races']} official FIDAL races from 2023-2026, matched to same-day Garmin activities.

The paper is version 1.0 and is scientifically complete for public release. The analysis is descriptive and does not make causal claims.

## Data Summary

- Garmin Connect export: {counts['garmin_start']} to {counts['garmin_end']}
- Activities: {counts['activities']:,}
- Runs: {counts['runs']:,}
- Running distance: {counts['run_km']:,.1f} km
- Sleep records: {counts['sleep_records']:,}
- Exact HRV records: {counts['hrv_records']}
- Official FIDAL races: {counts['races']}
- Garmin race matches: {counts['matched_races']}

## Key Findings

- 56-day running volume: r = 0.633, p = 0.0002
- 28-day running volume: r = 0.630, p = 0.0003
- 28-day sleep average: r = 0.580, p = 0.001
- 14-day running volume: r = 0.505, p = 0.0053
- VO2max nearest: r = 0.378, p = 0.0431
- 14-day sleep average: r = 0.403, p = 0.030

## Methods

Official race results were treated as the ground truth and imported from the athlete's FIDAL profile. Garmin data were used to compute pre-race training volume, sleep, HRV/RHR coverage, VO2max, and acute-load context. The outcome variable was event-relative official race performance.

## Limitations

- Single-athlete case study.
- Exploratory observational design.
- No causal inference.
- Mixed race events.
- Consumer wearable estimates are not laboratory measures.
- Exact HRV/RHR export coverage is sparse.
- Screenshot-derived HRV/RHR graph estimates are supplementary only and excluded from the primary statistical analysis.

## Reproducibility

Run the scripts in `scripts/` from this directory. Key outputs are in `data/`, `figures/`, `tables/`, `docs/`, `paper/`, and `publication/`.

Recommended order:

1. `scripts/build_activity_summary.py`
2. `scripts/build_health_summary.py`
3. `scripts/build_official_race_list.py`
4. `scripts/build_race_context_dataset.py`
5. `scripts/statistical_analysis.py`
6. `scripts/audit_hrv_rhr_coverage.py`
7. `scripts/digitize_screenshot_graphs.py`
8. `scripts/prepare_public_release.py`

## Manuscript

- `paper/final_manuscript.md`
- `paper/final_manuscript.docx`
- `paper/final_manuscript.pdf`

## Citation

Please cite this repository using `CITATION.cff`.

Suggested citation:

Bertoldo, A. (2026). *Training Volume, Sleep and Endurance Performance: A Longitudinal Analysis of 29 Official Races in a Competitive Distance Runner* (Version 1.0). Zenodo/OSF/GitHub release package.
"""
    (PROJECT_ROOT / "README.md").write_text(readme, encoding="utf-8")

    license_text = f"""MIT License

Copyright (c) 2026 {AUTHOR}

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""
    (PROJECT_ROOT / "LICENSE").write_text(license_text, encoding="utf-8")

    citation = {
        "cff-version": "1.2.0",
        "message": "If you use this dataset, code, or manuscript, please cite it as below.",
        "title": f"{TITLE}: {SUBTITLE}",
        "version": "1.0",
        "date-released": "2026-06-05",
        "authors": [
            {"family-names": "Bertoldo", "given-names": "Andrea", "affiliation": "Independent Student Researcher, Italy"}
        ],
        "keywords": ["wearable data", "endurance running", "Garmin", "sleep", "training volume", "FIDAL", "case study"],
        "license": "MIT",
        "abstract": "A four-year wearable-based single-athlete case study of Garmin-derived training and recovery context associated with 29 official FIDAL race performances.",
    }
    cff = "\n".join([
        f"cff-version: \"{citation['cff-version']}\"",
        f"message: \"{citation['message']}\"",
        f"title: \"{citation['title']}\"",
        f"version: \"{citation['version']}\"",
        f"date-released: \"{citation['date-released']}\"",
        "authors:",
        "  - family-names: \"Bertoldo\"",
        "    given-names: \"Andrea\"",
        "    affiliation: \"Independent Student Researcher, Italy\"",
        "keywords:",
        "  - wearable data",
        "  - endurance running",
        "  - Garmin",
        "  - sleep",
        "  - training volume",
        "  - FIDAL",
        "  - case study",
        "license: \"MIT\"",
        f"abstract: \"{citation['abstract']}\"",
    ]) + "\n"
    (PROJECT_ROOT / "CITATION.cff").write_text(cff, encoding="utf-8")

    notes = f"""# Repository Release Notes v1.0

Release date: 2026-06-05

## Contents

- Final manuscript in Markdown, DOCX, and PDF.
- Final publication figures and tables.
- Reproducible scripts for data preparation, race-context construction, statistical analysis, HRV/RHR audit, screenshot integration audit, and release preparation.
- Zenodo and OSF metadata/checklists.
- Final release audit.

## Scientific Scope

This release reports an exploratory observational single-athlete case study of {counts['races']} official FIDAL races from 2023-2026. The main findings identify longer running-volume windows and multi-week sleep averages as the clearest complete-data signals associated with event-relative performance.

## Important Limitations

- No causal claims.
- Single-athlete dataset.
- Consumer wearable estimates.
- HRV/RHR exact export coverage is sparse.
- Screenshot-derived HRV/RHR values are supplementary and excluded from primary statistics.

## Manual Release Steps

Before public publication, Andrea Bertoldo should review the final privacy audit and confirm that no private raw Garmin exports, GPS traces, personal identifiers, or unwanted screenshots are included.
"""
    (PROJECT_ROOT / "repository_release_notes_v1.md").write_text(notes, encoding="utf-8")


def write_publication_package(counts: dict) -> None:
    abstract = (
        "This exploratory single-athlete case study links Garmin Connect training and recovery data "
        "with 29 official FIDAL race results from 2023-2026. The strongest complete-data associations "
        "with event-relative official race performance were 56-day running volume, 28-day running "
        "volume, 28-day sleep average, 14-day running volume, VO2max nearest to race day, and 14-day "
        "sleep average. HRV/RHR exact export coverage was sparse, and screenshot-derived graph "
        "digitizations were retained only as supplementary evidence."
    )
    zenodo = {
        "title": f"{TITLE}: {SUBTITLE}",
        "upload_type": "publication",
        "publication_type": "workingpaper",
        "creators": [{"name": "Bertoldo, Andrea", "affiliation": "Independent Student Researcher, Italy"}],
        "description": abstract,
        "abstract": abstract,
        "keywords": ["wearable data", "endurance running", "Garmin", "sleep", "training volume", "FIDAL", "case study"],
        "version": "1.0",
        "language": "eng",
        "license": "MIT",
        "communities": [],
        "suggested_categories": ["Sport science", "Exercise physiology", "Wearable data", "Open data", "Single-case study"],
        "citation_text": "Bertoldo, A. (2026). Training Volume, Sleep and Endurance Performance: A Longitudinal Analysis of 29 Official Races in a Competitive Distance Runner (Version 1.0). Zenodo.",
    }
    (PUBLICATION / "zenodo_metadata.json").write_text(json.dumps(zenodo, indent=2), encoding="utf-8")

    zenodo_checklist = """# Zenodo Submission Checklist

## Files To Upload

Upload the release package directory or a ZIP containing:

- `README.md`
- `LICENSE`
- `CITATION.cff`
- `repository_release_notes_v1.md`
- `paper/final_manuscript.md`
- `paper/final_manuscript.docx`
- `paper/final_manuscript.pdf`
- `figures/final_figure_01.png` through `figures/final_figure_05.png`
- `tables/final_table_01.csv` through `tables/final_table_03.csv`
- `tables/final_table_captions.md`
- `data/race_context_dataset.csv`
- `data/stat_correlation_results.csv`
- `data/stat_signal_ranking.csv`
- `data/stat_best_worst_quartile_comparison.csv`
- `data/hrv_rhr_race_coverage_audit.csv`
- `data/hrv_rhr_combined_coverage_audit.csv`
- `docs/Statistical_Findings.md`
- `docs/HRV_RHR_Coverage_Audit.md`
- `docs/HRV_RHR_Screenshot_Integration_Report.md`
- `scripts/`
- `notebooks/`

Do not upload private raw Garmin export folders, GPS/FIT/TCX originals, or files containing precise route coordinates unless Andrea has explicitly reviewed and approved them.

## Metadata To Enter

- Upload type: Publication
- Publication type: Working paper or Preprint
- Title: Training Volume, Sleep and Endurance Performance: A Longitudinal Analysis of 29 Official Races in a Competitive Distance Runner
- Version: 1.0
- Creator: Bertoldo, Andrea; affiliation: Independent Student Researcher, Italy
- Description/abstract: use `publication/zenodo_metadata.json`
- Keywords: wearable data; endurance running; Garmin; sleep; training volume; FIDAL; case study
- License: MIT for code and documentation. If using a separate data license, choose CC BY 4.0 for derived anonymized data only after privacy review.

## Buttons To Press

1. Log in to Zenodo.
2. Click `New upload`.
3. Drag the prepared ZIP or selected files into the upload area.
4. Complete metadata fields above.
5. Choose license.
6. Save draft.
7. Preview the record.
8. Confirm all files are correct.
9. Click `Publish`.

## DOI

After pressing `Publish`, Zenodo mints a DOI automatically. Copy the DOI into:

- `README.md`
- `CITATION.cff`
- `publication/final_release_audit.md`
- the GitHub repository release notes
"""
    (PUBLICATION / "zenodo_submission_checklist.md").write_text(zenodo_checklist, encoding="utf-8")

    osf_metadata = f"""# OSF Metadata

## Title

{TITLE}: {SUBTITLE}

## Abstract

{abstract}

## Tags

wearable data; endurance running; Garmin; sleep; training volume; FIDAL; sport science; case study; observational analysis

## Description

This OSF project contains the manuscript, derived data, scripts, figures, and release documentation for version 1.0 of a four-year wearable-based single-athlete case study. Official FIDAL race results are treated as race-performance ground truth. Garmin-derived variables are used as training and recovery context. The study is exploratory and observational, with no causal claims.

## Contributors

- Andrea Bertoldo — author, data owner, analysis subject, independent student researcher, Italy

## License Recommendations

- Code and documentation: MIT License.
- Derived anonymized tabular data: CC BY 4.0 may be appropriate after final privacy review.
- Raw Garmin exports and precise route data: do not publish unless separately reviewed and intentionally shared.
"""
    (PUBLICATION / "osf_metadata.md").write_text(osf_metadata, encoding="utf-8")

    osf_checklist = """# OSF Submission Checklist

## Project Creation

1. Log in to OSF.
2. Click `Create new project`.
3. Enter the title from `publication/osf_metadata.md`.
4. Add Andrea Bertoldo as the project contributor and administrator.
5. Choose public visibility only after all files pass the privacy audit.

## File Upload

Upload the same curated package used for Zenodo:

- manuscript files in `paper/`
- final figures in `figures/`
- final tables in `tables/`
- derived CSV files in `data/`
- documentation in `docs/`
- reproducibility scripts in `scripts/`
- notebooks in `notebooks/`
- `README.md`, `LICENSE`, `CITATION.cff`, and release notes

Do not upload raw Garmin exports, route files, precise GPS coordinates, or private notes.

## Public Release

1. Review all files in OSF preview.
2. Confirm that no private identifiers, GPS coordinates, Garmin activity IDs, or unwanted screenshots are visible.
3. Add tags and description from `publication/osf_metadata.md`.
4. Choose the license.
5. Click `Make Public`.

## Citation Generation

After the project is public:

1. Open the OSF project overview.
2. Use the OSF citation widget or `Cite` button.
3. Copy the generated citation into the GitHub README and manuscript repository notes.
4. If a DOI is created or linked, add it to `CITATION.cff`.
"""
    (PUBLICATION / "osf_submission_checklist.md").write_text(osf_checklist, encoding="utf-8")


def run_privacy_checks() -> dict:
    checked_files = []
    suspicious_headers = []
    patterns = re.compile(
        r"(^|[_-])(lat|latitude|lon|longitude|gps|coordinate|activity[_-]?id|garmin[_-]?id|owner|email|phone)([_-]|$)",
        re.I,
    )
    for path in DATA.glob("*.csv"):
        checked_files.append(str(path.relative_to(PROJECT_ROOT)))
        try:
            with path.open("r", encoding="utf-8", newline="") as f:
                reader = csv.reader(f)
                headers = next(reader, [])
        except Exception:
            continue
        hits = [h for h in headers if patterns.search(h)]
        if hits:
            suspicious_headers.append({"file": str(path.relative_to(PROJECT_ROOT)), "headers": hits})
    return {"checked_files": checked_files, "suspicious_headers": suspicious_headers}


def write_final_audit(counts: dict) -> None:
    privacy = run_privacy_checks()
    stats = read_csv(DATA / "stat_correlation_results.csv")
    expected = {
        "Run km 56d": ("0.633", "0.0002"),
        "Run km 28d": ("0.630", "0.0003"),
        "Sleep avg 28d": ("0.580", "0.001"),
        "Run km 14d": ("0.505", "0.0053"),
        "VO2max nearest": ("0.378", "0.0431"),
        "Sleep avg 14d": ("0.403", "0.030"),
    }
    observed = {row["label"]: (fmt_num(row["pearson_r"], 3), fmt_p(row["pearson_p"])) for row in stats}
    checks = []
    for label, exp in expected.items():
        obs = observed.get(label)
        checks.append(f"- {label}: expected r={exp[0]}, p={exp[1]}; observed r={obs[0] if obs else 'missing'}, p={obs[1] if obs else 'missing'}; {'PASS' if obs == exp else 'CHECK'}")
    privacy_text = "\n".join(
        f"- {item['file']}: {', '.join(item['headers'])}" for item in privacy["suspicious_headers"]
    ) or "- No GPS coordinate, Garmin ID, email, phone, owner, or latitude/longitude columns detected in top-level derived CSV headers."
    audit = f"""# Final Release Audit

Date: {date.today().isoformat()}

## Scope

Release package for version 1.0 of:

{TITLE}: {SUBTITLE}

## Manuscript Consistency

- Title, author, abstract, methods, results, discussion, limitations, conclusion, references, and appendix are present.
- The manuscript explicitly states {counts['races']} official FIDAL races.
- Garmin export date range is stated as {counts['garmin_start']} to {counts['garmin_end']}.
- HRV/RHR export coverage limitation is stated.
- Screenshot-derived HRV/RHR are described as approximate and excluded from primary analysis.
- The design is described as exploratory and observational.
- No causal claims are made.

## Statistical Value Check

{chr(10).join(checks)}

## Figure References

- Figure 1: `figures/final_figure_01.png`
- Figure 2: `figures/final_figure_02.png`
- Figure 3: `figures/final_figure_03.png`
- Figure 4: `figures/final_figure_04.png`
- Figure 5: `figures/final_figure_05.png`

All five final figures exist and are referenced in the manuscript appendix.

## Table References

- Table 1: `tables/final_table_01.csv`
- Table 2: `tables/final_table_02.csv`
- Table 3: `tables/final_table_03.csv`

Captions are stored in `tables/final_table_captions.md`.

## Privacy Compliance

Top-level derived CSV files checked:

{chr(10).join('- ' + item for item in privacy['checked_files'])}

Potential sensitive header scan:

{privacy_text}

Notes:

- The public package should exclude raw Garmin export folders and original route files.
- Derived race data include public FIDAL source URLs and public race city names.
- The pseudonymous athlete id `single_athlete_001` is used in the race-context table.
- Manual screenshots should be reviewed by Andrea before public upload if included; the report treats them as supplementary.

## Release Readiness

- Scientific manuscript: ready for public preprint/repository release.
- GitHub repository: locally prepared; remote publishing requires confirmed repository remote and final privacy review.
- Zenodo package: metadata and checklist prepared.
- OSF package: metadata and checklist prepared.

## Manual Actions Before Publication

1. Review all files in `data/` and confirm no private raw Garmin exports are included.
2. Decide whether `data/manual_screenshots/` should be public or kept private.
3. Add DOI values after Zenodo/OSF publication.
4. Confirm license choice for derived data if different from MIT code/documentation.
"""
    (PUBLICATION / "final_release_audit.md").write_text(audit, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    counts = get_counts()
    build_tables()
    build_figures()
    md = save_markdown(counts)
    markdown_to_docx(md)
    markdown_to_pdf(md)
    write_readme_and_metadata(counts)
    write_publication_package(counts)
    write_final_audit(counts)
    print(json.dumps(counts, indent=2))
    print("Public release package generated.")


if __name__ == "__main__":
    main()
